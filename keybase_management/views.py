from os import EX_CANTCREAT

from django.contrib.auth.models import User
from numpy import RankWarning
from report_management.models import ReportsNotes
from django.db import IntegrityError, InternalError
from django.http import HttpResponse
from rest_framework.views import APIView
from rest_framework.viewsets import ModelViewSet
from rest_framework import viewsets
from keybase_management.serializers import KeybaseSerializer
from rest_framework.permissions import IsAuthenticated, DjangoModelPermissions
from rest_framework.authentication import TokenAuthentication
from .models import Keybase
from portfolio_management.views import perform_date_change
from rest_framework.response import Response
from rest_framework import status
from target_management.constants import TARGET_INDEX_RESOLVE, INDEX_LIST, KEYBASE_INDEX_LIST, TARGET_TYPE, \
    TARGET_SUB_TYPE
import datetime
from core_management.elasticsearch_handler_v2 import ElasticsearchHandler
from docx import Document
from rest_framework.decorators import action
from target_management.clean_data import clean_response
from target_management.data_cleaning import ElasticSearchResponseClean
from target_management.models import SocialTarget, KeybaseTarget, GenericTarget
from report_management.serializers import  NotesSerializer
import json 
from rest_framework.renderers import JSONRenderer
from core_management.models import Log
from core_management.serializers import LogSerializersversion2
# Create your views here.

es_object = ElasticsearchHandler()

es_clean = ElasticSearchResponseClean()


def get_attribute(target_type, sub_target_type):
    translate_target = dict(TARGET_TYPE)
    translate_subtarget = dict(TARGET_SUB_TYPE)
    target_type_translated = translate_target[str(target_type)]
    target_sub_type_translated = translate_subtarget[str(sub_target_type)]
    return target_type_translated, target_sub_type_translated


def get_index_for_es(target_type, target_sub_type, index):
    to_translate = str(target_type) + ',' + str(target_sub_type)
    translate_target = dict(TARGET_INDEX_RESOLVE)
    index_translated = translate_target[str(to_translate)]
    return_index = index_translated + index
    return return_index


def get_target_object(target_type=None, target_id=None, gtr=None):
    if gtr:
        type_gtr = gtr.split('_')[0]
        if type_gtr == 'st':
            specific_object = SocialTarget.objects.get(GTR=gtr)
            target_type = 'social'
        elif type_gtr == 'dt':
            specific_object = GenericTarget.objects.get(GTR=gtr)
            target_type = 'keybase'
        elif type_gtr == 'kb':
            specific_object = KeybaseTarget.objects.get(GTR=gtr)
            target_type = 'generic'
        else:
            specific_object = None

    elif target_type == 'social':
        if SocialTarget.objects.filter(id=target_id):
            specific_object = SocialTarget.objects.filter(id=target_id).first()
        else:
            specific_object = None

    elif target_type == 'keybase':
        if KeybaseTarget.objects.filter(id=target_id):
            specific_object = KeybaseTarget.objects.filter(id=target_id).first()
        else:
            specific_object = None

    elif target_type == 'generic':
        if GenericTarget.objects.filter(id=target_id):
            specific_object = GenericTarget.objects.filter(id=target_id).first()
        else:
            specific_object = None
    else:
        specific_object = None
    return specific_object, target_type
    return specific_object, target_type


class KeybaseViewSet(ModelViewSet):
    authentication_classes = (TokenAuthentication,)
    #queryset = Keybase.objects.all()
    pagination_class = None
    permission_classes = (IsAuthenticated, DjangoModelPermissions)

    def get_queryset(self):
        # return self.request.user.keybase_set.all().order_by('-created_on')
        return Keybase.objects.all().order_by('-created_on')
    def get_serializer_class(self):
        print(self.request.version)
        if self.request.version == 'v1':
            return KeybaseSerializer
        return KeybaseSerializer
    def list(self, request):
        auth_user = request.user.has_perm("keybase_management.view_keybase")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:
            if request.user.is_superuser:
                keybase=Keybase.objects.all().order_by('-created_on')
            else:
                keybase=Keybase.objects.filter(user=request.user).order_by('-created_on')
            keybase_serialized=KeybaseSerializer(keybase,many=True)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)
        response = {
            'message': "Successfully fetched",
            'result': keybase_serialized.data
        }
        return Response(response)

    def create(self, request):
        # auth_user = request.user.has_perm("keybase_management.can_create_keybase")
        # print(auth_user)
        # if auth_user == False:
        #     response = {
        #         'message': 'permission missing',
        #         'status': False,
        #         'result': None
        #         }
        #     return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:
            title = request.data.get('title', '')
            description = request.data.get('description', '')
            keywords = request.data.get('keywords', [])
            mentions = request.data.get('mentions', [])
            phrases = request.data.get('phrases', [])
            hashtags = request.data.get('hashtags', [])
            is_enabled = request.data.get('is_enabled', True)
            is_expired = request.data.get('is_expired', False)
            # created_on = request.data.get('created_on', '')
            # changed_created_date = datetime.datetime.fromtimestamp(float(created_on) / 1000)
            expire_on = request.data.get('expire_on', '')
            expire_on = datetime.datetime.fromtimestamp(float(expire_on) / 1000)
            updated_on = request.data.get('updated_on', '')

        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        try:
            Keybase.objects.create(title=title, description=description, keywords=keywords, mentions=mentions,
                                   phrases=phrases, hashtags=hashtags, is_enabled=is_enabled, is_expired=is_expired,
                                   expire_on=expire_on, updated_on=updated_on)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)
        response = {
            'message': "Keybase Created Successfully"
        }
        return Response(response)

    # def create(self, request, *args, **kwargs):
    #     # data = perform_date_change(request.data)
    #     # print('data_______', request.data)
    #     # print(data['keywords'])
    #     # if len(data['keywords']) > 10:
    #     #     response = {
    #     #         "message": "Keywords Limit Exceeded"
    #     #     }
    #     #     return Response(response, status=status.HTTP_406_NOT_ACCEPTABLE)
    #     # serializer = self.get_serializer(data=data)
    #     # serializer.is_valid(raise_exception=True)
    #     # if Keybase.get_keybase_count():
    #     #     try:
    #     #         serializer.save()
    #     #         headers = self.get_success_headers(serializer.data)
    #     #         response = {
    #     #             'message': 'Created Successfully',
    #     #             'status': True,
    #     #             'result': serializer.data
    #     #         }
    #     #         return Response(response, status=status.HTTP_201_CREATED, headers=headers)
    #     #
    #     #     except IntegrityError as ex:
    #     #         response = {
    #     #             'message': 'Integrity Error',
    #     #             'status': True,
    #     #             'result': str(ex)
    #     #         }
    #     #         return Response(response, status=status.HTTP_409_CONFLICT)
    #     #
    #     # else:
    #     #     response = {
    #     #         'message': 'Keybase Limit Exceeded For Today'
    #     #     }
    #     #     return Response(response)

    
    def keybase_details(self, request,**kwargs):
        try:
            data={}
            keybase_list=[]
            id=kwargs['id']
            # keybase_data=Keybase.objects.filter(id=id).first()
            # data['title']=keybase_data.title
            # data['description']=keybase_data.description
            # data['keywords']=keybase_data.keywords
            # data['mentions']=keybase_data.mentions
            # data['phrases']=keybase_data.phrases
            # data['hashtags']=keybase_data.hashtags
            # data['last_updated']=keybase_data.updated_on
            # data['keybase_user']=keybase_data.user.username
            # data['created_on']=keybase_data.created_on
            keybasetarget=KeybaseTarget.objects.filter(id=id).first()
            keybase_dict={}
            keybase_dict['GTR']=keybasetarget.GTR
            keybase_dict['keybase_target_created_by']=keybasetarget.user.username
            keybase_dict['keybase_target_created_on']=keybasetarget.created_on
            keybase_dict['keybase_title']=keybasetarget.keybase_title
            keybase_dict['updated_on']=keybasetarget.updated_on
            data['keybase_details']=keybase_dict
            reportnotes=Log.objects.filter(request_url='/v1/report/add/notes/')
            reportnotesserialized=LogSerializersversion2(reportnotes,many=True)
            report_json_data=JSONRenderer().render(reportnotesserialized.data)
            json_reportnotes_data = json.loads(report_json_data)
            newnote_list=[]
            for each_data in json_reportnotes_data:
                notedata={}
                try:
                    if each_data['request_data']['type']=='keybase' and  each_data['request_data']['report_id']==keybasetarget.GTR :
                        reportuser=User.objects.filter(id=each_data['request_username']).first()     
                        notedata['note_added_by']=reportuser.username
                        notedata['note_added_on']=each_data['request_time']
                        newnote_list.append(notedata)
                        # print('report_--------------------->genreated')                      
                    else:
                        pass
                except Exception as e:
                    # print(e)
                    pass
            data['note_data']=newnote_list 
            unique_note_contrubutors_list=[]
            dummyapi_data=Log.objects.filter(request_url='/v1/keybase/dummyapi/')
            serialized_dummyapi_data=LogSerializersversion2(dummyapi_data,many=True)
            dummyapi_json=JSONRenderer().render(serialized_dummyapi_data.data)
            json_dummy_api_data = json.loads(dummyapi_json)
            dummydata_list=[]
            for dummyapi in json_dummy_api_data:
                dummydata={}
                try:
                    if dummyapi['request_data']['id']==str(id):
                        dummy_api_user=User.objects.filter(id=dummyapi['request_username']).first()     
                        dummydata['note_added_by']=dummy_api_user.username
                        dummydata['note_added_on']=dummyapi['request_time']
                        dummydata_list.append(dummydata)
                        # print('report_--------------------->dummmyy api')                      
                    else:
                        pass
                except Exception as e:
                    pass
            unique_dummy_contributors=[]        
            for data_dum in dummydata_list:
                unique_dummy_contributors.append(data_dum['note_added_by'])        
            unique_contributors_report=set(unique_dummy_contributors) 
            data['report_generated_data']=dummydata_list
            for each_note in newnote_list:
                unique_note_contrubutors_list.append(each_note['note_added_by'])
            unique_noteadd_contrubutors=set(unique_note_contrubutors_list)
            data['total_contributors']=1+len(unique_noteadd_contrubutors)+len(unique_contributors_report)       

        except Exception as e:
            print('-----------------',e)
            response={
                'error':e
            }
            
        response = {
            'message': data
        }
        return Response(response)
    def dummyapi(self,request,**kwargs):
        try:
            id=request.data['id']
        except Exception as error:
            response={
                'message':str(error)
            }     
            return Response(response)

        response={
                'result':id
            }     
        return Response(response)    

    # def keybase_available_key(self,request,**kwargs):
    #         try:
    #             gtr=kwargs['gtr']
    #             keybase_status=[]
    #             available_key=False
    #             print("------------gtr-----------",gtr)
    #             keybase_target_marked=KeybaseTarget.objects.filter(GTR=gtr).first()
    #             keybase_status=keybase_target_marked.target_status_string
    #             keybase_status_split=keybase_status.split(',')
    #             print('------------->',keybase_status_split)
    #             if keybase_status_split[-1]=='Successful! data store to BDS ':
    #                 available_key=True
    #             else:
    #                 available_key=False    
    #         except Exception as error:
    #             response={
    #                 'message':str(error)
    #             }     
    #             return Response(response)

    #         response={
    #                 'result':available_key
    #             }     
    #         return Response(response)


    def update(self, request, *args, **kwargs):
        try:
            # data = perform_date_change(request.data)
            data = request.data
            id = data['id']
            title = data['title']
            description = request.data['description']
            hashtags = request.data['hashtags']
            phrases = request.data['phrases']
            keywords = request.data['keywords']
            mentions = request.data['mentions']
            expire_on = request.data['expire_on']
            # created_on = request.data['created_on']
            is_enabled = request.data['is_enabled']
            is_expired = request.data['is_expired']
            changed_expire_date = datetime.datetime.fromtimestamp(float(expire_on) / 1000)
            # changed_created_date = datetime.datetime.fromtimestamp(float(created_on) / 1000)
            if len(data['keywords']) > 10:
                response = {
                    "message": "Keywords Limit Exceeded"
                }
                return Response(response)
            if data.get('id', None):
                keybase_data = Keybase.objects.filter(id=id).values()
                keybase_data.update(title=title, description=description, hashtags=hashtags, phrases=phrases,
                                    keywords=keywords, mentions=mentions,
                                    expire_on=changed_expire_date, is_enabled=is_enabled,
                                    is_expired=is_expired)
            response = {
                'message': 'Successfully Updated'
            }
            return Response(response)

        except Exception as error:
            print(error)

        response = {
            'message': 'error'
        }
        return Response(response)


    def keybase_filter(self, request):
        try:

            start_date = request.data['start_date']
            end_date = request.data['end_date']
            keybase = Keybase.objects.filter(created_on__range=[start_date, end_date])
            serialized_data = KeybaseSerializer(keybase, many=True)

        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        response = {
            'result': serialized_data.data
        }
        return Response(response)    


# class KeybaseReport(viewsets.ViewSet):
#     authentication_classes = (TokenAuthentication, )
#     permission_classes = (IsAuthenticated,)
#
#     @action(detail=True, methods=['POST'])   

          
def KeybaseReport(request,*args, **kwargs):
    auth_user=request.user.has_perm("keybase_management.can_generate_report") 
    if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)  
    id = kwargs['id']
    limit = kwargs['limit']
    # data_dict = request.data.copy()
    GTR = kwargs['GTR']
    # if 'GTR' in data_dict:
    #     gtr, limit = data_dict['GTR'], data_dict['limit']
    specific_object, target_type = get_target_object(gtr=GTR)
    if specific_object:
        gtr = specific_object.GTR
        target_type_es, target_sub_type_es = get_attribute(target_type=specific_object.target_type,
                                                           sub_target_type=specific_object.target_sub_type)
    document = Document()
    responses = dict()
    keybase_id = KeybaseTarget.objects.filter(id=id).values()
    for key in keybase_id:
        key_id = key['keybase_id']
    keybase_obj = Keybase.objects.filter(id=key_id).values()
    document.add_heading('Keybase Report', 0)
    document.add_heading('keybase Details', 1)
    try:
        for obj in keybase_obj:
            keybase_id = obj['id']
            keybase_title = obj['title']
            keywords = obj['keywords']
            hashtags = obj['hashtags']
            keybase_creation = obj['created_on']
    except Exception as error:
        response = {'message': str(error)}
        return Response(response)
    if keybase_obj:
        document.add_paragraph('ID: ' + str(keybase_id))
        document.add_paragraph('Title: ' + str(keybase_title))
        document.add_paragraph('Keywords: ' + str(keywords))
        document.add_paragraph('Hashtags: ' + str(hashtags))
        document.add_paragraph('Created On: ' + str(keybase_creation))

    index_list = KEYBASE_INDEX_LIST
    for index in index_list:

        index_es = get_index_for_es(target_type_es, target_sub_type_es, str(index))
        res_list, count = es_object.get_response_by_gtr(index_es, gtr, limit)
        responses.update({index: res_list})
        # print(res_list)

    cleaned_responses, mined_data = clean_response(responses)

    if cleaned_responses:
        try:
            categorization, sentiments, emotions = es_clean.graph_data(cleaned_responses)
            cleaned_responses.update({'profile_categorization_count': categorization})
            cleaned_responses.update({'profile_sentiments_count': sentiments})
            cleaned_responses.update({'profile_emotions_count': emotions})
        except:
            cleaned_responses.update({'profile_categorization_count': {}})
            cleaned_responses.update({'profile_sentiments_count': {}})
            cleaned_responses.update({'profile_emotions_count': {}})
        try:
            cleaned_word_cloud = es_clean.clean_response_cloud(mined_data)
            cleaned_responses.update({'word_cloud': cleaned_word_cloud['word_cloud']})
        except:
            cleaned_responses.update({'word_cloud': ""})

        try:
            cleaned_most_common_words = es_clean.clean_most_common_words(mined_data)
            cleaned_responses.update({'most_common_words': cleaned_most_common_words['common']})
        except:
            cleaned_responses.update({'most_common_words': []})
        try:
            cleaned_most_common_words = es_clean.clean_most_common_hashtags(mined_data)
            cleaned_responses.update({'most_used_hashtags': cleaned_most_common_words['hash_tags']})
        except:
            cleaned_responses.update({'most_used_hashtags': []})
        try:
            cleaned_most_common_words = es_clean.clean_response_summary(mined_data)
            cleaned_responses.update({'target_summary': cleaned_most_common_words['target_summary']})
        except:
            cleaned_responses.update({'target_summary': []})
        try:
            cleaned_most_common_words = es_clean.clean_response_freq(mined_data)
            cleaned_responses.update(
                {'post_freq_graph': cleaned_most_common_words['graph']})
        except:
            cleaned_responses.update({'post_freq_graph': []})
    else:
        pass
    result = cleaned_responses

    # #################################### FACEBOOK ###################################

    document.add_heading('Keybase Result Counts', 1)
    facebook_count = []
    document.add_heading('Facebook', 1)
    # Users
    count = 1
    facbook_users = []
    for data in result['facebook_users']:
        facbook_users.append((data['full_name'], data['username']))

    if len(facbook_users) > 0:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Full Name'
        hdr_cells[0].text = 'Username'

        for full_name, username in facbook_users:
            row_cells = table.add_row().cells
            row_cells[0].text = full_name
            row_cells[1].text = username

    else:
        document.add_paragraph('No Data')


    # Posts

    facebook_posts = []
    document.add_heading('Facebook Posts', 2)
    for data in result['facebook_posts']:
        facebook_posts.append((data['post_text'], str(data['reactions']), str(data['comments']),
                               data['timestamp']))

    if len(facebook_posts) > 0:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Post Text'
        hdr_cells[1].text = 'Reactions'
        hdr_cells[2].text = 'Comments'
        hdr_cells[3].text = 'Date'

        for post_text, reaction, comments, date in facebook_posts:
            row_cells = table.add_row().cells
            row_cells[0].text = post_text
            row_cells[1].text = reaction
            row_cells[2].text = comments
            row_cells[3].text = date
    else:
        document.add_paragraph('No Data')


    # Pages
    facebook_pages = []
    document.add_heading('Facebook Pages', 2)
    for data in result['facebook_pages']:
        facebook_pages.append((data['full_name'], data['username']))

    if len(facebook_pages) > 0:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Full Name'
        hdr_cells[1].text = 'Username'

        for full_name, username in facebook_pages:
            row_cells = table.add_row().cells
            row_cells[0].text = full_name
            row_cells[1].text = username
    else:
        document.add_paragraph('No Data')

    # Groups
    document.add_heading('Facebook Groups', 2)
    facebook_group = []
    for data in result['facebook_groups']:
        facebook_group.append((data['full_name'], data['username']))

    if len(facebook_group) > 0:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Full Name'
        hdr_cells[1].text = 'Username'

        for full_name, username in facebook_group:
            row_cells = table.add_row().cells
            row_cells[0].text = full_name
            row_cells[1].text = username
    else:
        document.add_paragraph('No Data')
    # ################################### INSTAGRAM ######################################

    document.add_heading('Instagram', 1)

    # Users
    instagram_user = []
    document.add_heading('Users', 2)
    for data in result['instagram_users']:
        instagram_user.append((data['name'], data['username']))

    if len(instagram_user) > 0:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Username'
        for name, username in instagram_user:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = username

    else:
        document.add_paragraph('No Data')


    # Hashtags
    document.add_heading('Hashtags', 2)
    instagram_hashtags = []
    for data in result['instagram_hashtags']:
        instagram_hashtags.append((data['name'], data['post_count']))

    if len(instagram_hashtags) > 0:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Post Count'

        for name, post_count in instagram_hashtags:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = str(post_count)
    else:
        document.add_paragraph('No Data')

    # ################################### TWITTER ##############################################

    document.add_heading('Twitter', 1)

    # Users
    twitter_users = []
    for data in result['twitter_tweets']:
        for d in data['users']:
            twitter_users.append((d['description'], d['name'], d['username']))

    if len(twitter_users) > 0:
        table = document.add_table(rows=1, cols=3)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Description'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Username'

        for description, name, username in twitter_users:
            row_cells = table.add_row().cells
            row_cells[0].text = description
            row_cells[1].text = name
            row_cells[2].text = username
    else:
        document.add_paragraph('No Data')

    document.add_heading('Youtube')

    document.add_heading('videos', 2)
    youtube_videos = []
    for data in result['videos']:
        youtube_videos.append((data['title']))

    if len(youtube_videos) > 0:
        table = document.add_table(rows=1, cols=1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Title'

        for title in youtube_videos:
            row_cells = table.add_row().cells
            row_cells[0].text = title
    else:
        document.add_paragraph('No Data')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Excutive_summary.docx'

    document.save(response)
    return response
