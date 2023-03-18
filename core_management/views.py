
import datetime
from sys import implementation
from target_management.target_identify_socket import TargetIdentifySocialMedia

from rest_framework.views import exception_handler
from keybase_management.models import Keybase
from case_management.models import Case
from portfolio_management.models import Individual, Group, Event
from django.contrib.auth.models import User, Group as gp
from django.db.models import Q, query
from django.http import JsonResponse, HttpResponse
from django_eventstream import send_event, get_current_event_id
from rest_framework.authentication import TokenAuthentication
from core_management.constants import COUNTRY_LIST, CHANNEL_LIST, TWITTER_COUNTRIES, GOOGLE_COUNTRIES
from case_management.views import perform_upload ,perform_upload_frs
from target_management.models import SocialTarget, KeybaseTarget, GenericTarget
import time
from docx import Document
from .models import Notification
from .serializers import *
from rest_framework import viewsets, status
from rest_framework.decorators import action, parser_classes
from django_currentuser.middleware import get_current_user, get_current_authenticated_user
from rest_framework import authentication
from rest_framework.response import Response
from core_management import elasticsearch_handler
from target_management import ess_controller
from target_management import ais_controller
from portfolio_management.models import LinkedData
from target_management import bds_controller
from bs4 import BeautifulSoup
from rest_framework import permissions
from . import elasticsearch_handler_v2
from .models import FakeIdentity, IPLogger, Log, Log_history
from .models import FakeIdentity, IPLogger, NewsMonitoring, AutoML, ImageAnalysisFRS
from OCS_Rest.dev import HDFS_ROOT_PATH2,HDFS_ROOT_PATH
from core_management.hdfs_client import send_file
from case_management.serializers import CaseDashboardsSerializer
from keybase_management.serializers import KeybaseSerializer
from target_management.serializers import SocialTargetListSerializer, GenericTargetListSerializer, KeybaseTargetListSerializer
from django.contrib.admin.models import LogEntry
import requests
import json
from .log_filter import log_filter_url
from account_management.models import Team, Team_members
from .sftp_upload import FTP_Client
from .spark_automl import SparkApiController
import random
from target_management.signals import get_attribute
import pandas as pd
from django.core.paginator import Paginator
from target_management.models import UserNews
from .text_processing import process_data
from .kill import killjob
from target_management.constants import TARGET_INDEX_RESOLVE, INDEX_LIST, KEYBASE_INDEX_LIST, GENERIC_INDEX_LIST, \
    INDEX_PLATFORM_ALL

# from rest_framework import filters
# from django_filters.rest_framework import DjangoFilterBackend
# import json


es_object = elasticsearch_handler_v2.ElasticsearchHandler()
ess_object = ess_controller.EssApiController()
ais_object = ais_controller.AisApiController()
bds_object = bds_controller.BDSController()


class NewsView(viewsets.ViewSet):
    """
    Viewset for data processing tools
    """
    authentication_classes = (authentication.TokenAuthentication,)

    @action(detail=True, methods=['POST'])
    def top_news_crawl(self, request):

        auth_user = request.user.has_perm("core_management.news_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        de_serialize = NewsCrawelerSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        # news_responses = obj.get_response_by_type(index='general_data', document_type='news', channel='ABP',
        #                                          limit=10000)
        if set(keys) == {'channel', 'limit'}:
            channel, limit = data_dict['channel'], data_dict['limit']
            try:
                news_responses = es_object.get_response_by_index(index='general_data_news',
                                                                 channel=str(channel), limit=limit)
                if len(news_responses) ==0:
                    response = {'message': 'News Not found',
                            'status': True,
                            'result': []}
                    return Response(response)                                          
                response = sorted(news_responses, key = lambda i: i['created_on'],reverse=True)
                # print(response[0]['priority'])
                
                priority_dict = {
                    "highest": 4,
                    "high": 3,
                    "medium": 2,
                    "low": 1,
                }
                
                response = sorted(response, key = lambda i: priority_dict[i['priority']],reverse=True)

                for each_resp in response:
                    list_of_words = each_resp['word_cloud'].split(' ')
                    unique_values = set(list_of_words)
                    each_resp.update({'word_cloud': [{'tag':i,"count":list_of_words.count(i)} for i in unique_values]})

                # BBC 
                # if news_responses:
                #     result = news_responses
                #     news_date = datetime.datetime.fromtimestamp(int(result[0]['created_on'] / 1000))
                #     current_datetime = datetime.datetime.now() - datetime.timedelta(hours=0, minutes=20)
                #     print(news_date)
                #     if news_date < current_datetime:
                #         send_event('notification', 'message',
                #                    {"Notification_message": "No latest news from {}".format(result[0]['Channel'])})
                #         response = {'message': 'News Response',
                #                     'Alert': "No latest news from {}".format(result[0]['Channel']),
                #                     'status': True,
                #                     'result': result}
                #     else:
                #         response = {'message': 'News Response',
                #                     'status': True,
                #                     'result': result}
                #
                #     return Response(response, status=status.HTTP_200_OK)
                #
                # else:
                # result = None
                response = {'message': 'News Response',
                            'status': True,
                            'result': response}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def top_news(self, request):
        auth_user = request.user.has_perm("core_management.news_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        de_serialize = NewsSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'top', 'news_site', 'channel_url'}:
            top, news_site, channel_url = data_dict['top'], data_dict['news_site'], data_dict['channel_url']
            try:
                result = ess_object.news_crawling(top=top, news_site=str(news_site), channel_link=channel_url)
                response = {'message': 'News',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def get_news_search(self, request):
        auth_user = request.user.has_perm("core_management.news_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        de_serialize = NewsSearchSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'query'}:
            query = data_dict['query']
            try:
                result = ess_object.get_news_search(news=str(query))
                response = {'message': 'News',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)


class ToolsView(viewsets.ViewSet):
    """
    Viewset for data processing tools
    """
    authentication_classes = (authentication.TokenAuthentication,)

    @action(detail=True, methods=['POST'])
    def identity(self, request):
        auth_user = request.user.has_perm("avatar_management.amu_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        de_serialize = ToolsIdentifySerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'nationality', 'gender', 'age'}:
            nationality, gender, age = data_dict['nationality'], data_dict['gender'], data_dict['age']
            try:
                result = ess_object.fake_identity_generator(nationality=nationality, gender=gender, age=age)

                FakeIdentity.objects.create(data=result, created_by=request.user)
                # if FakeIdentity:
                # full_name = result['data']['personal_info']['full_name']
                # username = result['data']['online_info']['username']
                # password = result['data']['online_info']['password']
                # age = result['data']['personal_info']['age']
                # gender = result['data']['personal_info']['gender']
                # email = result['data']['online_info']['email_address']

                # FakeIdentity.objects.create(full_name=full_name, username=username, age=age, email=email,
                #                             gender=gender, password=password)

                response = {'message': 'Target Status',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    def all_fake_identities(self, request):
        try:
            if request.user.is_superuser:
                obj = FakeIdentity.objects.all().values()
            else:
                obj = FakeIdentity.objects.filter(created_by = request.user).values()

            response = {
                'results': obj
            }
            return Response(response)
        except Exception as error:
            response = {
                'Error': str(error)
            }
            return Response(response)

    def delete_fake_identity(self, request, *args, **kwargs):
        try:
            id = kwargs['id']

            obj = FakeIdentity.objects.filter(id=id)

            obj.delete()

            response = {
                'message': 'Deleted Successfully'
            }
            return Response(response)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

    @action(detail=True, methods=['POST'])
    def dark(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        de_serialize = ToolsDarkSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'query'}:
            query = data_dict['query']
            try:
                result = ess_object.dark_web_search(query=query)
                response = {'message': 'Target Status',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
        return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def scrappers(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        de_serialize = ToolsScrapperSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'type', 'query'}:
            query = data_dict['query']
            scrapper_type = data_dict['type']
            try:
                if scrapper_type == 'scholar':
                    result = ess_object.google_scholar_data_scraper(query=query)
                elif scrapper_type == 'patent':
                    result = ess_object.google_patents_data_scraper(query=query)
                elif scrapper_type == 'amazon':
                    result = ess_object.amazon_data_scraper(query=query)
                elif scrapper_type == 'daraz':
                    result = ess_object.daraz_data_scraper(query=query)
                else:
                    result = None
                response = {'message': 'Scrapper Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
        return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def text(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        de_serialize = ToolsTextSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        print("Request By User: ", user)
        data_dict = request.data.copy()
        print('$$$$$$$$$$$$$$$$$$4')
        print(request.data)
        keys = [*data_dict]
        if set(keys) == {'type', 'query'}:
            query = data_dict['query']
            scrapper_type = data_dict['type']

            # ********************************************************************************
            # *********************************** Old Code ***********************************
            # ********************************************************************************
            #     target_object = LinkedData.objects.create(query=query, type=scrapper_type)
            #     if scrapper_type == 'sentiment':
            #         result = bds_object.text_processing_req(input_string=query,
            #                                                 algorithm_type='sentiment', data_id=target_object.id)
            #     elif scrapper_type == 'hashtag':
            #         result = bds_object.text_processing_req(input_string=query,
            #                                                 algorithm_type='most_used_hashtags',
            #                                                 data_id=target_object.id)
            #     elif scrapper_type == 'cloud':
            #         result = bds_object.text_processing_req(input_string=query,
            #                                                 algorithm_type='word_clouds', data_id=target_object.id)
            #     elif scrapper_type == 'word':
            #         result = bds_object.text_processing_req(input_string=query,
            #                                                 algorithm_type='common_words', data_id=target_object.id)
            #     else:
            #         result = None
            #     print("Before the sleep statement")
            #     print(result)
            #     data = None
            #     count_flag = 0
            #     time.sleep(10)
            #     while data is None:
            #         time.sleep(5)
            #         print("After the sleep statement")
            #         if LinkedData.objects.filter(id=target_object.id):
            #             linked_data = LinkedData.objects.filter(id=target_object.id).first()
            #             data = linked_data.data
            #             if len(data) <= 0:
            #                 data = None
            #         else:
            #             data = None
            #         if count_flag >= 22:
            #             break
            #         else:
            #             count_flag += 1

            # ********************************************************************************
            # *********************************** New Code ***********************************
            # ********************************************************************************                result = process_data(input_str=query, algorithm_type=scrapper_type)
            try:
                type_dict = {
                    'word': 'common_words',
                    'sentiment': 'sentiment',
                    'hashtag': 'most_used_hashtags',
                    'cloud': 'word_clouds'
                }
                
                data = process_data(input_str=query, algorithm_type=type_dict[scrapper_type])
                
                response = {'message': 'Text Processing Result',
                            'status': True,
                            'result': data}
                return Response(response, status=status.HTTP_200_OK)
            except Exception as E:
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_204_NO_CONTENT)
        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
        return Response(response, status=status.HTTP_204_NO_CONTENT)


class TwitterTools(viewsets.ViewSet):
    """
    Viewset for twitter tools
    """
    authentication_classes = (authentication.TokenAuthentication,)

    def tweets_search(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:
            category = request.data['tweets_category']
            # scroll = int(request.data['scroll'])
            all_of_these_words = request.data['all_of_these_words']
            exact_phrase = request.data['exact_phrase']
            any_of_these_words = request.data['any_of_these_words']
            none_of_these_words = request.data['none_of_these_words']
            hashtags = request.data['hashtags']
            language = request.data['language']
            from_these_accounts = request.data['from_these_accounts']
            to_these_accounts = request.data['to_these_accounts']
            mentioning_these_accounts = request.data['mentioning_these_accounts']
            replies = request.data['replies']
            only_replies = request.data['only_replies']
            links = request.data['links']
            only_tweet_with_links = request.data['only_tweet_with_links']
            min_replies = request.data['min_replies']
            min_likes = request.data['min_likes']
            min_retweets = request.data['min_retweets']
            from_date = request.data['from_date']
            to_date = request.data['to_date']

            tweets_responses = ess_object.tweets_search_result(category, all_of_these_words, exact_phrase,
                                                               any_of_these_words, none_of_these_words, hashtags,
                                                               language, from_these_accounts, to_these_accounts,
                                                               mentioning_these_accounts, replies, only_replies,
                                                               links, only_tweet_with_links, min_replies, min_likes,
                                                               min_retweets, from_date, to_date)

            return Response(tweets_responses, status=status.HTTP_200_OK)

        except Exception as error:
            print("ERROR", error)

        response = {
            'message': 'Request Field missing'
        }
        return Response(response, status=status.HTTP_200_OK)

    @action(detail=True, methods=['POST'])
    def phrase_near_location(self, request):
        user = request.user
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'phrase', 'location'}:
            try:
                result = ess_object.tweets_near_location(query=data_dict['phrase'],
                                                         location=data_dict['location'])
                response = {'message': 'Twitter search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def near_location_within_miles(self, request):
        user = request.user
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'distance', 'location'}:
            try:
                result = ess_object.tweets_near_location_within_miles(location=data_dict['location'],
                                                                      distance=data_dict['distance'])
                response = {'message': 'Twitter search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def phrase_search(self, request):
        user = request.user
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'query'}:
            try:
                result = ess_object.tweets(query=str(data_dict['query']))
                response = {'message': 'Twitter search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)


class LookupTools(viewsets.ViewSet):
    """
    Viewset for twitter tools
    """
    authentication_classes = (authentication.TokenAuthentication,)

    @action(detail=True, methods=['POST'])
    def ip_short(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        de_serialize = LookupIpSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        print("Request By User: ", user)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'title', 'description', 'url'}:
            try:
                result = ess_object.create_payload(title=data_dict['title'],url=data_dict['url'])
                response = {'message': 'Ip payload',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def image_reverse_lookup(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print('request', request)
        print('request_data', request.data)
        print('FILES', request.FILES)
        print("Request By User: ", user)
        print("Request By User: ", user)
        de_serialize = ImageSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'image_url', 'image'}:
            try:
                data = perform_upload(request.data, 'image')
                print(data,'-------------------------------------------------')
                result = TargetIdentifySocialMedia().start_targetidentify(target="image_reverse_lookup",query=data['image_url'])
                # result = ess_object.image_reverse_lookup(image=None,
                #                                          url=data['image_url'])
                response = {'message': 'Image reverse lookup',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_200_OK)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def domain_ip_info(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        data_dict = request.data.copy()
        de_serialize = IpSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        keys = [*data_dict]
        if set(keys) == {'ip'}:
            try:
                result = ess_object.get_domains_ip_info(domian=str(data_dict['ip']))
                response = {'message': 'ip search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def domain_info(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        user = request.user
        print("Request By User: ", user)
        de_serialize = DomainSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        keys = [*data_dict]
        if set(keys) == {'domain'}:
            try:
                result = ess_object.get_domains_info(domian=str(data_dict['domain']))
                response = {'message': 'domain search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)

    @action(detail=True, methods=['POST'])
    def tag_response(self, request):
        user = request.user
        print("Request By User: ", user)
        de_serialize = TagResponseSerializer(data=request.data)
        de_serialize.is_valid(raise_exception=True)
        data_dict = request.data.copy()
        print('data', data_dict)
        keys = [*data_dict]
        if set(keys) == {'tag'}:
            try:
                result = ess_object.get_tag_originator(tagName=str(data_dict['tag']))
                response = {'message': 'tag search Result',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            except Exception as E:
                print(E)
                response = {'message': 'Error while Getting Response',
                            'status': True,
                            'result': str(E)}
                return Response(response, status=status.HTTP_404_NOT_FOUND)

        else:
            response = {
                'message': 'Data Keys Missing ',
                'status': False,
                'result': None
            }
            return Response(response, status=status.HTTP_204_NO_CONTENT)


def latest_trend(created_date):
    print("creates dataa")
    trend_date = datetime.datetime.fromtimestamp(int(created_date / 1000))
    current_datetime = datetime.datetime.now() - datetime.timedelta(hours=0, minutes=20)
    print(trend_date, "trend date", current_datetime)
    if current_datetime > trend_date:
        print(" old trend  ")
        return False
    else:
        print("latest ")
        return True


class BigView(viewsets.ViewSet):
    """
    View sets of Big view
    """
    authentication_classes = (authentication.TokenAuthentication,)

    @action(detail=True, methods=['GET'])
    def trends_crawl(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            ess_object.youtube_trends()
            ess_object.reddit_trends()
            # for google_country in TWITTER_COUNTRIES:
            #     ess_object.google_trends(country=google_country)
            # for twitter_country in GOOGLE_COUNTRIES:
            #     ess_object.twitter_trends(country=twitter_country)

            ess_object.twitter_trends(country='WORLD')

            for twitter_country in TWITTER_COUNTRIES:
                ess_object.twitter_trends(country=twitter_country)

            for google_country in GOOGLE_COUNTRIES:
                ess_object.google_trends(country=google_country)

            result = "Marked"
            response = {'message': 'Trends',
                        'status': True,
                        'result': result}
            return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['GET'])
    def news_crawl(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            # for channel in CHANNEL_LIST:
            #     ess_object.news_crawling(news_site=channel, top=30)
            channels = UserNews.objects.values_list("news_name", "channel_url").distinct()
            
            for channel in channels:
                ess_object.news_crawling(news_site=channel[0], top=30, channel_link=channel[1])

            result = "Marked"
            response = {'message': 'News',
                        'status': True,
                        'result': result}
            return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['POST'])
    def add_target_twitter_worldwide(self, request):
        user = request.user
        print("Request By User: ", user)
        if 'country' in request.data:
            country = request.data['country'].lower().replace(" ","-")
            if country not in TWITTER_COUNTRIES:
                response = {'message': 'Country not found in twitter trends OCS',
                            'status': True,
                            'result': []}
                return Response(response, status=status.HTTP_200_OK)

        else:
            country = 'pakistan'
        try:
            result = es_object.get_twitter_trends(index='trends_twitter_country_trends',
                                                  created_on_gte=datetime.datetime.timestamp
                                                  (datetime.datetime.now() + datetime.timedelta(hours=-48*2)),
                                                  created_on_lte=datetime.datetime.timestamp
                                                  (datetime.datetime.now() + datetime.timedelta(hours=24)), limit=200,
                                                  country=country)
            if len(result) ==0:
                response = {'message': 'no result found',
                            'status': True,
                            'result': result}
                return Response(response)
            print(result)
            get_latest_date = latest_trend(result[0]['created_on'])
            if get_latest_date:
                response = {'message': 'twitter trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                send_event('notification', 'message', {"Notification_message": "Twitter trends are not latest"})
                response = {'message': 'twitter trends',
                            "Alert": "Twitter trends are not latest",
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['GET'])
    def reddit_trends(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            result = es_object.get_trends(index='trends_reddit_trends',
                                          created_on_gte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=-48*2)),
                                          created_on_lte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=24)),
                                          limit=15)
            if len(result)== 0:
                response = {'message': 'NO Reddit data trends found',
                            'status': True,
                            'result': result}
                            
                return Response(response, status=status.HTTP_200_OK)
            get_latest_date = latest_trend(result[0]['created_on'])
            if get_latest_date:
                response = {'message': 'Reddit trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                send_event('notification', 'message', {"Notification_message": 'Reddit trends are not latest'})
                response = {'message': 'Reddit trends',
                            'Alert': 'Reddit trends are not latest',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['GET'])
    def func_get_ml_model(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            result = es_object.get_ml_model(index='ml_models')
            get_data = result
            if get_data:
                response = {'message': 'Auto Ml data ',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                response = {'message': 'No data',
                            'Alert': 'No data ',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)
        

    @action(detail=True, methods=['GET'])
    def youtube_trends(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            
            result = es_object.get_trends(index='trends_youtube_trends',
                                          created_on_gte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=-48*2)),
                                          created_on_lte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=24)), limit=15)
            print(result)                              

            get_latest_date = latest_trend(result[0]['created_on'])
            
            if get_latest_date:
                response = {'message': 'Youtube trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                response = {'message': 'Youtube trends',
                            "Alert": "Youtube trends are not latest",
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['GET'])
    def twitter_world_trends(self, request):
        user = request.user
        print("Request By User: ", user)
        try:
            result = es_object.get_trends(index='trends_twitter_country_trends',
                                          created_on_gte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=-48*2)),
                                          created_on_lte=datetime.datetime.timestamp
                                          (datetime.datetime.now() + datetime.timedelta(hours=24)),
                                          limit=15)
            print(result)
            get_latest_date = latest_trend(result[0]['created_on'])
            if get_latest_date:
                response = {'message': 'twitter trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                response = {'message': 'twitter trends',
                            "Alert": "Twitter trends are not latest",
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['POST'])
    def google_world_trends(self, request):
        user = request.user
        print("Request By User: ", user)
        if 'country' in request.data:
            country = request.data['country'].lower().replace(" ","_")
            if country not in GOOGLE_COUNTRIES:
                response = {'message': 'Country not found in OCS',
                            'status': True,
                            'result': []}
                return Response(response, status=status.HTTP_200_OK)

        else:
            country = 'argentina'
        try:
            if country == "all":

                result = es_object.get_google_trends(index='trends_google_trends_all',
                                                 limit=10000, country=country)
            else:
                result = es_object.get_google_trends_by_country(index='trends_google_trends_all',
                                               limit=10000, country=country)
            if len(result) == 0:
                response = {'message': 'google trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

            get_latest_date = latest_trend(result[0]['created_on'])
            if get_latest_date:
                response = {'message': 'google trends',
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)
            else:
                response = {'message': 'google trends',
                            "Alert": "google trends are not latest",
                            'status': True,
                            'result': result}
                return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)

    @action(detail=True, methods=['GET'])
    def covid_data_get(self, request):
        user = request.user
        print("Request By User: ", user)
        url_to_get = 'https://covid.amcharts.com/data/js/total_timeline.js'
        try:
            import json
            import requests
            r = requests.get(url_to_get)
            content_data = r.content
            soup = str(BeautifulSoup(content_data)).split('var covid_total_timeline=[')[1]
            soup.encode('unicode_escape')
            res = soup.replace('\\', '')
            # content_data = soup.decode_contents()
            # print(json.loads(content_data_conv))
            response = {'message': 'covid',
                        'status': True,
                        'result': [res]}
            return Response(response, status=status.HTTP_200_OK)

        except Exception as E:
            print(E)
            response = {'message': 'Error while Getting Response',
                        'status': True,
                        'result': str(E)}
            return Response(response, status=status.HTTP_404_NOT_FOUND)


class RecentTarget(viewsets.ModelViewSet):
    queryset = SocialTarget.objects.all()
    authentication_classes = (TokenAuthentication,)

    def list(self, request, *args, **kwargs):
        instagram_total_count = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='in')).count()
        linkedin_total_coutn = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='ln')).count()
        redit_total_count = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='rd')).count()
        twitter_total_count = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='tw')).count()
        facebook_total_count = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='fb')).count()
        youtube_total_count = SocialTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='yt')).count()
        keybase_total_count = KeybaseTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='kb')).count()
        generic_total_count = GenericTarget.objects.filter(Q(
            created_on__icontains=datetime.datetime.date(datetime.datetime.now()), target_type='gn')).count()
        custom_total_count = keybase_total_count + generic_total_count
        return JsonResponse(
            {"instagram_total_count": [{"name": "Instagram", "total_count": instagram_total_count}],
             "linkedin_total_coutn": [{"name": "LinkedIn", "total_count": linkedin_total_coutn}],
             "redit_total_count": [{"name": "Redit", "total_count": redit_total_count}],
             "twitter_total_count": [{"name": "Twitter", "total_count": twitter_total_count}],
             "facebook_total_count": [{"name": "Facebook", "total_count": facebook_total_count}],
             "youtube_total_count": [{"name": "Youtube", "total_count": youtube_total_count}],
             "custom_total_count": [{"name": "custom", "total_count": custom_total_count}]})


class Targets(viewsets.ModelViewSet):
    authentication_classes = (TokenAuthentication,)

    def list(self, request, *args, **kwargs):
        social_targets = SocialTarget.objects.all().count()
        keybase_targets = KeybaseTarget.objects.all().count()
        generic_targets = GenericTarget.objects.all().count()
        targets_added = social_targets + keybase_targets + generic_targets
        expired_social = SocialTarget.objects.filter(expire_on__lte=datetime.datetime.now()).count()
        expired_keybase = KeybaseTarget.objects.filter(expire_on__lte=datetime.datetime.now()).count()
        expired_generic = GenericTarget.objects.filter(expire_on__lte=datetime.datetime.now()).count()
        expired_targets = expired_social + expired_generic + expired_keybase
        social_periodic_targets = SocialTarget.objects.filter(periodic_interval__gt=0).count()
        keybase_periodic_targets = KeybaseTarget.objects.filter(periodic_interval__gt=0).count()
        generic_periodic_targets = GenericTarget.objects.filter(periodic_interval__gt=0).count()
        periodic_targets = social_periodic_targets + generic_periodic_targets + keybase_periodic_targets
        facebook_count = SocialTarget.objects.filter(target_type='fb').count()
        twitter_count = SocialTarget.objects.filter(target_type='tw').count()
        instagram_count = SocialTarget.objects.filter(target_type='in').count()
        youtube_count = SocialTarget.objects.filter(target_type='yt').count()
        reddit_count = SocialTarget.objects.filter(target_type='rd').count()
        linkedin_count = SocialTarget.objects.filter(target_type='ln').count()
        word_cloud = SocialTarget.objects.values_list('full_name', flat=True)
        final_word_cloud = list(word_cloud)
        dates = []

        all_social_targets = SocialTarget.objects.all()
        for targets in all_social_targets:
            dates.append(targets.created_on.date())

        return JsonResponse({"Targets_Added": [{
            "category": "Social Targets", "value": social_targets
        }, {"category": "keybase Targets", "value": keybase_targets},
            {"category": "Total Targets Count", "value": targets_added}, {"category": "Total Expired Targets",
                                                                          "value": expired_targets},
            {"category": "Total Periodic Targets", "value": periodic_targets}],

            "Social_Media_Count": [{"Category": "Facebook", "value": facebook_count}, {"Category": "Twitter",
                                                                                       "value": twitter_count},
                                   {"Category": "Youtube", "value": youtube_count}, {"Category": "Instagram",
                                                                                     "value": instagram_count},
                                   {"Category": "Linkedin", "value": linkedin_count}, {"Category": "Reddit",
                                                                                       "value": reddit_count}],
            "Word_Cloud": final_word_cloud,
            "dates": dates
        })


class CountAll(viewsets.ModelViewSet):
    authentication_classes = (TokenAuthentication,)

    def list(self, request, *args, **kwargs):
        keybase_count = Keybase.objects.all().count()
        user_count = User.objects.all().count()
        case_count = Case.objects.all().count()
        individual_portfolio_count = Individual.objects.all().count()
        group_portfolio_count = Group.objects.all().count()
        event_portfolio_count = Event.objects.all().count()
        keybase_word_cloud = list(Keybase.objects.values_list('title', flat=True))
        case_word_cloud = list(Case.objects.values_list('case_title', flat=True))
        portfolio_word_cloud = list(Individual.objects.values_list('title', flat=True))
        portfolio_word_cloud.extend(list(Group.objects.values_list('title', flat=True)))
        portfolio_word_cloud.extend(list(Event.objects.values_list('title', flat=True)))

        return JsonResponse({"Count": [{
            "Category": "No. of Users", "value": user_count
        }, {
            "Category": "No. of Keybase", "value": keybase_count
        }, {
            "Category": "No. of Cases", "value": case_count
        }, {
            "Category": "No. of Individual Portfolios", "value": individual_portfolio_count
        }, {
            "category": "No. of Group Portfolios", "value": group_portfolio_count
        }, {
            "category": "No. of Event Portfolios", "value": event_portfolio_count
        }],
            "Keybase_word_cloud": keybase_word_cloud,
            "Case_word_cloud": case_word_cloud,
            "Portfolio_word_cloud": portfolio_word_cloud
        })


def send_notification(*args, **kwargs):
    """
    function will send the notification to user

    @params <Required>
    """
    try:
        event_id = get_current_event_id(['notification'])

        print(int(str(event_id).split(':')[1]))
        print('doneee')
        print(" Sending Notifications ", kwargs['messege'])
        user = User.objects.get(id=kwargs['user'])
        notify = Notification(user=user, Notification_message=kwargs['messege'], read_status=False)
        notify.save()
        print(notify.id)
        send_event('notification', 'message', {"Notification_message": kwargs['messege']})
        print("Sent")

    except Exception as e:
        print("-------------------", e, "--------------------------")


class NotificationViewSet(viewsets.ViewSet):
    """
    Notification Viewset to display unread notifications
    """
    authentication_classes = (TokenAuthentication,)

    def list(self, request):
        print(get_current_event_id(['notification']))
        last_id = get_current_event_id(['notification'])
        data = Notification.objects.filter(user=request.user, read_status=False).order_by('-id')
        serilizer = NotificationSerializer(data, many=True)
        return Response({"result": serilizer.data, "last_id": last_id})




class AllNotification(viewsets.ViewSet):
    """
    Notification to display all notifications of specific user
    """

    authentication_classes = (TokenAuthentication,)

    def get_queryset(self):
        return Notification.objects.filter(user=self.request.user).order_by('-id')

    def list(self, request):
        data = self.get_queryset()
        serilizer = NotificationSerializer(data, many=True)
        return Response({"result": serilizer.data})

    def update(self, request, *args, **kwargs):
        notification = Notification.objects.get(id=kwargs['pk'])
        notification.read_status = True
        notification.save()
        return Response({"result": " Notification read status"})


class Rapideye(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)

    def rapideye_response(self, request):
        final_list = []
        try:
            platform = request.data['platform']
            attributes = request.data['attributes'].split(',')
            values = request.data['values'].split(',')
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)
        try:
            final_response = es_object.rapideye_fetched_data(platform, attributes, values)
            if platform == "youtube" or platform == "linkedin":
                for res in final_response:
                    if res[attributes[0]].lower() == values[0].lower():
                        final_list.append(res)
            else:
                data_of_response = {}
                for each_response in final_response:
                    if each_response["username"] not in data_of_response:
                        data_of_response[each_response["username"]] = each_response
                    else:
                        latest_gtr = each_response["GTR"].split("_")[-1]
                        old_gtr = data_of_response[each_response["username"]]["GTR"].split("_")[-1]
                        if latest_gtr> old_gtr:
                            data_of_response[each_response["username"]] = each_response



                
                for key, res in data_of_response.items():
                    if res[attributes[0]].lower() == values[0].lower():
                        final_list.append(res)
        except Exception as error:
            response = {
                'message': "Error while getting response from Elasticsearch",
                'result': str(error)
            }
            return Response(response)

        response = {
            'result': final_list
        }
        return Response(response)

    def response_count(self, request):
        try:
            platform = request.data['platform']
        except Exception as error:
            response = {
                'error': str(error)
            }
            return Response(response)
        try:
            total_count = es_object.es_response_count(platform)
        except Exception as error:
            response = {
                'error': str(error)
            }
            return Response(response)

        response = {
            'total_count': total_count
        }
        return Response(response)


class IP_Logger(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)

    def ip_logger(self, request):
        code = request.data['code']
        print(code)
        try:
            ess_response = ess_object.track_ip(code)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)
        try:
            city = ess_response['data']['IP_data']['city']
            country = ess_response['data']['IP_data']['country']
            domain_name = ess_response['data']['IP_data']['domain_name']
            isp = ess_response['data']['IP_data']['isp']
            region = ess_response['data']['IP_data']['region']

            for res in ess_response['data']['data']:
                ip_add = res['ip_addr']
                time = res['timestamp']
                user_agent = res['user_agent']
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        try:
            print("SAVING DATA")
            IPLogger.objects.create(city=city, country=country, domain_name=domain_name, isp=isp, region=region,
                                    ip_add=ip_add, time=time, user_agent=user_agent)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        response = {
            'result': 'Done'
        }
        return Response(response)


    def get_ip_logger(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:
            data = requests.get("http://majidahmed.pythonanywhere.com/all_ip")
            # obj = IPLogger.objects.all()
            # ser = IPLoggerSerializer(obj, many=True)
            response = {
                'result': json.loads(data.text),

            }
            return Response(response)

        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

    def delete_ip_logger(self, request, *args, **kwargs):
        id = kwargs['id']

        try:
            response = requests.post("http://majidahmed.pythonanywhere.com/delete_ip",data={'id':id})
            print(response)
            # obj = IPLogger.objects.filter(id=id)
            # print()
            # obj.delete()
            response = {
                'message': 'Delete Successfully'
            }
            return Response(response)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)


class CrawlersDetail(viewsets.ViewSet):

    authentication_classes = (TokenAuthentication,)

    def crawler_status(self, request):
        auth_user = request.user.has_perm("core_management.history_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:

            response = ess_object.get_crawler_status()

            return Response(response)

        except Exception as error:
            respone = {
                'message': str(error)
            }
        return Response(respone)

    def terminate_task(self, request):
        try:
            id = request.data['id']

            ess_response = ess_object.get_task_terminate(id)

            response = {
                'result': ess_response
            }
            return Response(response)

        except Exception as error:
            response = {
                'Error': str(error)
            }
            return Response(response)


class DataRetention(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)

    def all_targets(self, request):
        try:
            start_date = request.data['start_date']
            end_date = request.data['end_date']
            start_date = datetime.datetime.fromtimestamp(float(start_date) / 1000)
            end_date = datetime.datetime.fromtimestamp(float(end_date) / 1000)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)
        try:
            social_obj = SocialTarget.objects.filter(created_on__range=[start_date, end_date]).values()
            generic_obj = GenericTarget.objects.filter(created_on__range=[start_date, end_date]).values()
            keybase_obj = KeybaseTarget.objects.filter(created_on__range=[start_date, end_date]).values()
        except Exception as error:
            response = {
                'Error': str(error)
            }
            return Response(response)
        response = {
            'Social_Targets': social_obj,
            'Generic_Targets': generic_obj,
            'Keybase_Targets': keybase_obj
        }
        return Response(response)

    def delete_by_gtr(self, request):
        try:
            gtr = request.data['gtr']

            response = es_object.delete_by_gtr(gtr)
            obj = SocialTarget.objects.get(GTR=gtr)
            obj.delete()

            if response:
                response = {
                    'message': 'Deleted Successfully'
                }
                return Response(response)

        except Exception as error:
            response = {
                'Error': str(error)
            }
            return Response(response)
    
    def update_categorization(self, request):
        try:
            GTR = request.data['gtr']
            post_id = request.data['post_id']
            new_categorization = request.data['new_categorization']

            response1 = es_object.update_categorization(GTR, post_id, new_categorization)

            # response = {
            # 'update_categorization': response,
            # }

        
            response = {
                'message': "Successfully update categorization",
                'success':True,
                'response1':response1
                }
         

            return Response(response)


        except Exception as error:
            response = {
                'message': "Error update categorization",
                'Error': str(error),
                'success':False
            }
            return Response(response)
    

    def update_sentiment(self, request):
        try:
            GTR = request.data['gtr']
            post_id = request.data['post_id']
            sentiment = request.data['sentiment']

            sentiment_lower=[]
            for senti in sentiment:
                sentiment_lower.append(senti.lower())


            if sentiment_lower[0] == 'positive' or sentiment_lower[0] == 'negative' or sentiment_lower[0] == 'neutral':
                response1 = es_object.update_sentiment(GTR, post_id, sentiment_lower)
                response = {
                'message': "Successfully update sentiment",
                'success':True,
                'data':response1
                }
            else:
                response = {
                'message': "UnSuccessfully update sentiment",
                'success':False
                }

            return Response(response)


        except Exception as error:
            response = {
                'message': "Error update categorization",
                'Error': str(error),
                'success':False
            }
            return Response(response)

    def update_additional_post_comment(self, request):
       
        try:
            
            gtr = request.data['gtr']
            post_id = request.data['post_id']
            comment = request.data['comment']
            commenter_name = request.data['commenter_name']
            # print(request.user)
            user_name = str(request.user)
            
            response = es_object.update_additional_post_comment(gtr, post_id, comment,commenter_name)

            # response = {
            # 'update_additional_post_comment': response,
            # }
            print(response)
        
            response = {
                'message': 'successfully added',
                'success':True
                }
           
            return Response(response)

        except Exception as error:
            response = {
                'message': "Error in update post comment",
                'Error': str(error),
                'success':False
            }
            return Response(response)

    def get_twitter_top_ten_user_view(self, request):
        try:
            GTR = request.data['gtr']
            
            response1 = es_object.get_twitter_top_ten_user(GTR)
            # response1 = es_object.get_top_ten_domain_name(GTR)

            response = {
                'message': "Successfully update post comment",
                'success':True,
                'response1':response1
                }
           
            return Response(response)

        except Exception as error:
            response = {
                'message': "error",
                'Error': str(error),
                'success':False
            }
            return Response(response)



            


class UserNotifications(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)
    # permission_classes = (permissions.AllowAny,)

    def all_user_notifications(self, request):
        notifications = {}
        user = User.objects.get(username=request.user)
        social_obj = SocialTarget.objects.filter(user_id=user.id).order_by('-created_on').values_list('user_names', 'target_status_string', 'created_on')[:10]
        keybase_obj = KeybaseTarget.objects.filter(user_id=user.id).order_by('-created_on').values_list('keybase_title','target_status_string', 'created_on')[:10]
        generic_obj = GenericTarget.objects.filter(user_id=user.id).order_by('-created_on').values_list('title', 'target_status_string', 'created_on')[:10]

        notifications['social'] = social_obj
        notifications['keybase'] = keybase_obj
        notifications['generic'] = generic_obj

        return Response(notifications)


class NewsMonitioringView(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)

    def news_monitioring(self, request):
        try:
            news_name = request.data['news_name']
            NewsMonitoring.objects.create(user=request.user, news=news_name)
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)


class ModelTraining(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)

    def model_training(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:
            file = request.data['file']
            hdfs_app_directory = 'training_data'
            file_name = "{0}.csv".format('training_data')
            send_file(dir_to_store="{0}/{1}/".format(HDFS_ROOT_PATH, hdfs_app_directory), data_to_store=file,
                     case_name=file_name)

            AutoML.objects.create(user=request.user, file_name=file_name)

        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        response = {
            'message': 'File Uploaded Successfully'
        }
        return Response(response)
    
    def get_image_analysis_FRS(self, request ):
        try:
            data = []
            title = self.request.query_params.get('title')
            if title:
                images = ImageAnalysisFRS.objects.filter(title = title)
                if images:
                    image_data = es_object.get_image_analysis_FRS(images[0].title)
                    for each_data in image_data:
                        try:
                            each_data["target_id"] = SocialTarget.objects.get(GTR=each_data['GTR']).id
                        except:
                            pass
                        print(each_data["GTR"])
                    data.append({"title":images[0].title, "image_data":image_data})
                    return Response({"results":data, "message":"data found"})
                else:
                    return Response({"results":data, "message":"No data found"})

            images = ImageAnalysisFRS.objects.all()
            for title in images:
                image_data = es_object.get_image_analysis_FRS(images[0].title)
                for each_data in image_data:
                    try:
                        each_data["target_id"] = SocialTarget.objects.get(GTR=each_data['GTR']).id
                    except:
                        pass
                data.append({"title":title.title, "image_data":image_data})
            # data.append(es_object.get_image_analysis_FRS("test1"))
            return Response({"results":data, "error":"no errors"})
        except Exception as e:
            print(e)
            return Response({"results":[],"errors":e})


    def get_all_ml_models(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        try:

            automl = AutoML.objects.all().order_by("-upload_on")
            serializer = AutoMLSerializer(automl, many=True)
            
            final_data = []
            
            for model in serializer.data:
                try:
                    if model['model_details']['subtype']:
                        print('----------------------------')
                except Exception as e:
                    final_data.append(model)
            return Response(final_data)
            # response1 = es_object.get_all_ml_models()
            # if not response1 or len(response1):
            #     response = {
            #         'message': "UnSuccessfully update sentiment",
            #         'success':False,
            #         "response1":response1
            #     }
            # return Response(response)


        except Exception as error:
            response = {
                'message': "Error update categorization",
                'Error': str(error),
                'success':False
            }
            return Response(response)
    def closeassociates(self,request,**kwargs):
        GTR=kwargs['GTR']
        CTR=kwargs['CTR']
        target_type=kwargs['target_type']
        target_subtype=kwargs['target_subtype']

        try:
            print('in try')
            response=es_object.close_associates(GTR ,CTR , target_type, target_subtype)
        except Exception as error:
            print('------------->in except')
            response={
                'message': str(error)
            }
                        
        return Response(response)  
    def delete_frs(self,request,**kwargs):
        title=kwargs['title']

        try:
            print('in try')
            response=es_object.delete_image_analysis_FRS(title)
        except Exception as error:
            print('------------->in except')
            response={
                'message': str(error)
            }
                        
        return Response(print('-------------------------------------->',response))   
    def get_ml_model(self, request):
        try:
            model_name = request.data["model_name"]

            automl = AutoML.objects.filter(model_name = model_name).first()
            if automl.status != "DONE":
                response = {
                        'message': automl.message,
                        "status": automl.status,
                        'success':True,
                        "response1":{}
                    }
                return Response(response)
            else:
                response1 = es_object.get_ml_model(model_name,automl.type)
                if not response1 or len(response1):
                    response = {
                        'message': automl.message,
                        'status':automl.status,
                        'success':True,
                        "response1":response1
                    }
                return Response(response)
        except Exception as error:
            response = {
                'message': "Get AutoML Error",
                'Error': str(error),
                'success':False
            }
            print(error)
            return Response(response)

    def delete_ml_model(self, request):
        try:
            model_name = request.data["model_name"]

            automl = AutoML.objects.filter(model_name = model_name).first()
            
            if automl:
                # delete from postgres
                automl.delete()
            
                # delete from ES
                es_object.delete_ml_model(index='ml_models', name=model_name)
            
                response = {
                    'message': 'Deleted Successfully!',
                    'model': automl.model_name,
                    'success': True
                }
            else:
                response = {
                    "message": "Auto-ML model is not present"
                }            

            return Response(response)
        except Exception as error:
            response = {
                'message': "Get AutoML Error",
                'Error': str(error),
                'success': False
            }

            return Response(response)

    def cancel_ml_model(self, request):
        try:            
            model_name = request.data["model_name"]

            automl = AutoML.objects.filter(model_name = model_name).first()
            print(automl)
            
            if automl:
                # update  from postgres
                # this needs implementation
                # automl.update()
            
                try:
                    # kill process
                    # killjob()
                    pass
                except:
                    pass

                response = {
                    'message': 'Job killed successfully!',
                    'model': automl.model_name,
                    'success': True
                }
            else:
                response = {
                    "message": "Auto-ML model is not present"
                }            

            return Response(response)
        except Exception as error:
            response = {
                'message': "Get AutoML Error",
                'Error': str(error),
                'success': False
            }

            return Response(response)
    
    def deploy_model(self,request):
        model_name = request.data['model_name']
        print(model_name)
        if not model_name:
            return Response({
                'message': "Model name is mandatory",
                'Error': "401",
                'success':False
            })
        
        response1 = es_object.deploy_model(model_name)
        return Response({"message":"Successfully deploy model","response":response1})

    def discard_model(self,request):
        model_name = request.data['model_name']
        print(model_name)
        if not model_name:
            return Response({
                'message': "Model name is mandatory",
                'Error': "401",
                'success':False
            })
        
        response1 = es_object.discard_model(model_name)
        return Response({"message": "Successfully discard model", "response": response1})


    def training_keybase(self, request):

        # serializer = TrainingKeywordsSerializer(data = request.data)
        # if serializer.is_valid():
        #     data = serializer.data
        # print(request.data['training_time'])
        # print(request.data['is_csv_file'])
        # print(request.data['csv_name'])
        # print(request.data['is_datawarehouse'])
        # print(request.data['training_from'])
        # print(request.data['training_to'])
        
        print("data------>training", request.data)
        try:
            training_time = request.data['training_time']
        except:
            training_time = 3600


        try:
            type_ = request.data['type']
        except:
            type_ = "CATEGORIZATION"

        is_csv_file = True
        csv_name = request.data['csv_name']
        file_name = "no_file"

        is_datawarehouse_checker = request.data['is_datawarehouse']

        if is_datawarehouse_checker=="False" or is_datawarehouse_checker=="false":
            is_datawarehouse= False
        elif is_datawarehouse_checker=="True" or is_datawarehouse_checker=="true":
            is_datawarehouse = True


        convert_to_epochs = lambda x:  datetime.datetime.timestamp(datetime.datetime.strptime(x, "%Y-%m-%d")) * 1000
        try:
            if is_datawarehouse == False:
                training_from = None
                training_to = None
            else:

                training_from = convert_to_epochs(request.data['training_from'])
                training_to = convert_to_epochs(request.data['training_to'])
                print(training_from,training_to)
        except :
            training_from = None
            training_to = None
        try:
            file = request.data['file']
        except:
            file = 'null'
            file_name = "no_file"
        print("********************")
        print(file)
        print(type(file))
        print('********************')
        # try:
        #     if file == 'null':
        #         # print("yeeeeee")
        #         file_check == False
        #     else:
        #         print("noooooooooooo")
        #         print(file)
        # except:
        #     print("pata nhi")
        #     print(file)
        # if AutoML.objects.filter(status= "IN_PROGRESS").exists():
        #     return Response({"message":"Job is ready in progress"})

        try:
            try: 
                if file:
                    data = pd.read_csv(file).copy()
                    # print(data)
                    print(len(data))
                    print(data.shape)
                    print(data.columns[0])
                    # check_rows_count = pd.read_csv(file, nrows=400).copy().shape[0]
                    if data.shape[0] > 5000:
                        return Response({"message":"Rows should be less then '5000' Provided '{0}' rows data".format(data.shape[0])})
                    if data.shape[1] >= 2:
                        print("found data")
                        if data.columns[0] != 'text' or data.columns[1] != 'label':
                            return Response({"message":"first column head should be 'text' and Secound 'label'"})
                        if data.columns[0] == 'text' and data.columns[1] == 'label':
                            n = random.randint(0,random.randint(1, 2000))
                            # file = request.data['file'] 
                            file_name = csv_name+str(n) # place random integer
                            hdfs_app_directory = 'data/{}'.format(type_.lower())
                            file_name = "{0}.csv".format(file_name)
                            print("file name ", file_name)
                            send_file(dir_to_store="{0}/{1}/".format(HDFS_ROOT_PATH2, hdfs_app_directory), data_to_store=file,
                                    case_name=file_name)
                            
                            model_name = "{}_{}".format(type_,str(round(datetime.datetime.now().timestamp())))

                            AutoML.objects.create(user=request.user, file_name=file_name, model_name = model_name, type = type_)
                            sparkml = SparkApiController()
                            check = sparkml.ml_training(training_time, is_csv_file,file_name, is_datawarehouse, training_from, training_to,model_name,type_)

                            #print("check---------", check.content)
                            # print(check.status_code)
                            return Response({"message":"Traning start Successfully",
                                                "check":check})
                        else:
                            return Response({"message":"file '{0}' column name should be 'text' and label".format(request.data['csv_name'])})
                    else:
                        return Response({"message":"file '{0}' have no data".format(request.data['csv_name'])})
            except:
                if file == 'null':
                    is_csv_file = False
                    model_name = "{}_{}".format(type_,str(round(datetime.datetime.now().timestamp())))
                    AutoML.objects.create(user=request.user, file_name=file_name, model_name = model_name, type = type_)
                    sparkml = SparkApiController()
                    check = sparkml.ml_training(training_time, is_csv_file ,csv_name, is_datawarehouse, training_from, training_to,model_name,type_)
                    print("check---------", check.content)
                    # print(check.status_code)
                    return Response({"message":"Traning start Successfully",
                                        "check":check})
        except Exception as e:
            return Response({"message":"file '{0}' uploaded".format(e)})

            
        # if data.empty():
        #     print("yes found data")
        # sparkml = SparkApiController()
        
        # check = sparkml.ml_training(data['training_time'],data['is_csv_file'],data['csv_name'],data['is_datawarehouse'],data['training_from'],data['training_to'])
        # AutoML.objects.create(user=request.user, file_name=request.data['csv_name'])
        # return Response({"message":"file '{0}' uploaded".format(request.data['csv_name'])})
        # return Response({"status":check,"result":serializer.data})

        # return Response({"errer":serializer.errors})


    # def get_all_ml_models(self,request):
    #     query = {
    #         "query":{
    #             "bool":{
    #             "must":[
    #                 {
    #                     "match_all":{}
    #                 }
    #             ]
    #             }
    #         }
    #     }
        


    # def get_one_ml_model(self,request):
    #     model_name = request.get['model_name']
    #     query = {
    #         "query":{
    #             "bool":{
    #             "must":[
    #                 {
    #                     "match":{
    #                         "name":model_name
    #                     }
    #                 }
    #             ]
    #             }
    #         },
    #         "size":1
    #     }



    def file_upload_for_training(self, request):
        try:

            # serializer = TrainingFileUploadSerializer(data = request.data)
            # if serializer.is_valid():
            #     return Response(serializer.data)
            # return Response(serializer.errors)
            
            n = random.randint(0,random.randint(1, 2000))
            file = request.data['file'] 
            file_name = request.data['file_name'] 
            hdfs_app_directory = 'data'
            file_name = "{0}{1}.csv".format(file_name,n)
            send_file(dir_to_store="{0}/{1}/".format(HDFS_ROOT_PATH2, hdfs_app_directory), data_to_store=file,
                     case_name=file_name)
            AutoML.objects.create(user=request.user, file_name=file_name)
            return Response({"message":"file '{0}' uploaded".format(file_name)})
        except Exception as error:
            response = {
                'message': str(error)
            }
            return Response(response)

        # response = {
        #     'message': 'File Uploaded Successfully'
        # }
        # return Response(response)


    def automl_notify(self, request):

        model_name = request.data['model_name'] 
        message = request.data['message'] 
        status = request.data['status']

        try:
            automl = AutoML.objects.filter(model_name = model_name).first()

            automl.message = message
            automl.status = status
            automl.save()

            return Response({"message":"Updated Successfully"})
        except Exception as e:
            return Response({"message":"error on update",
                            "error":str(e)})


class Report(viewsets.ModelViewSet):

    permission_classes = (permissions.AllowAny,)
    @action(detail=True, metods=["get"], permission_classes=((permissions.AllowAny,)))
    def report_generation(self,request,*args, **kwargs):
        # auth_user =request.user.has_perm("core_management.tools_permission")
        # print(auth_user)
        # if auth_user == False:
        #         response = {
        #             'message': 'permission missing',
        #             'status': False,
        #             'result': None
        #             }
        #         return Response(response,status=status.HTTP_403_FORBIDDEN)

        
        from_date = kwargs['from_date']
        to_date = kwargs['to_date']
        category = kwargs['category'].split(',')
        social_media_platform = kwargs['social_media'].split(',')
        action = kwargs['action']
        # print("---------------------------", category, "------------------------------------")
        # res = get_categorization_visual("1611687600000", ["GTR", "created_on", "posts.url_c", "url",
        #                                 "categorization.predictions", "categorization.confidence"])

        output_formate = ["GTR", "created_on", "posts.url_c", "url", "categorization.predictions",
                        "categorization.confidence"]

        response = es_object.get_categorization_visual(from_date, to_date, output_formate)
        # response = json.loads(response)
        all_social_medias = ['facebook', 'twitter', 'instagram', 'linkedin', 'reddit']
        all_pakistan_news = ['dawn', 'arynews', 'aaj', 'geo', 'duniya', 'gnn', 'hum', 'indus', 'ktn', 'khyber', 'neo', 'samaa',
                            'bol', 'apna', 'abbtak', 'express', 'waqt', 'such', '24news', '92news']
        world_wide_news = ['bbc', 'fox', 'cnn', 'sky', 'msnbc', 'aljazeera', 'uronews', 'alarabiya', 'ndtv', 'cgtn',
                        'wion', 'cnc', 'i24', 'atn', 'sahar', 'fox', 'saam', 'abp', 'mewsmax', 'zee']
        social_media = []
        world_news = []
        paksitan_news = []
        blogs = []

        document = Document()
        try:

            for res in response:
                target_type = res['target_type']
                gtr = res['fields']['GTR'][0]
                if target_type == 'keybase':
                    try:
                        keybase_target_obj = KeybaseTarget.objects.get(GTR=gtr)
                        username = keybase_target_obj.user_names
                    except:
                        print("           keybase             ")

                elif target_type == 'generic_crawler':
                    try:
                        generic_target_obj = GenericTarget.objects.get(GTR=gtr)
                        username = generic_target_obj.title
                        main = 'Title'
                    except:
                        print("           generic_crawler             ")

                else:
                    try:
                        social_target_obj = SocialTarget.objects.get(GTR=gtr)
                        username = social_target_obj.user_names
                        main = 'Username'
                    except :
                        print("           Social_crawler             ")

                categorization = res['fields']['categorization.predictions'][0]
                if categorization in category:

                    if 'posts.url_c' in res['fields']:
                        post_urls = res['fields']['posts.url_c'][0]
                    if 'url' in res['fields']:
                        post_urls = res['fields']['url'][0]
                    created_on = datetime.datetime.fromtimestamp(float(res['fields']['created_on'][0]) / 1000)
                    confidence = float(res['fields']['categorization.confidence'][0]) * 100

                    count = 0
                    for i in all_social_medias:
                        if i in post_urls:
                            count = count + 1
                            social_media.append((username, target_type, post_urls, created_on, categorization, confidence))

                    for i in all_pakistan_news:
                        if i in post_urls:
                            count = count + 1
                            paksitan_news.append((username, target_type, post_urls, created_on, categorization, confidence))

                    for i in world_wide_news:
                        if i in world_news:
                            count = count + 1
                            world_news.append((username, target_type, post_urls, created_on, categorization, confidence))

                    if count == 0:
                        blogs.append((username, target_type, post_urls, created_on, categorization, confidence))
        except Exception as error:
            pass
            print("error______ report ", error)
        try:

            if action == 'data':
                final_social_media = []
                all_data = {}
                for i in social_media:
                    if i[1] in social_media_platform:
                        final_social_media.append(i)
                all_data['social_media'] = final_social_media
                all_data['pakistan_news'] = paksitan_news
                all_data['world_news'] = world_news
                all_data['blogs'] = blogs
                return JsonResponse(all_data)
        except:
            message = "no data"
            return JsonResponse(message)

        # ---------------------------------- Social Media ------------------------------------
        if action == 'report':

            if len(social_media) > 0:
                document.add_heading('Social Media')
                table = document.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = main
                hdr_cells[1].text = 'Target Type'
                hdr_cells[2].text = 'URL'
                hdr_cells[3].text = 'Date'
                hdr_cells[4].text = 'Categorization'
                hdr_cells[5].text = 'Confidence in %'

                for username, target_type, created_on, post_url, categorization, confidence in social_media:
                    if target_type in social_media_platform:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(username)
                        row_cells[1].text = str(target_type)
                        row_cells[2].text = str(created_on)
                        row_cells[3].text = str(post_url)
                        row_cells[4].text = str(categorization)
                        row_cells[5].text = str(confidence)

            # -------------------------------- Pakistan News ----------------------------------

            if len(paksitan_news) > 0:
                document.add_heading('Pakistan News')
                table = document.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = main
                hdr_cells[1].text = 'URL'
                hdr_cells[2].text = 'Target Type'
                hdr_cells[3].text = 'Date'
                hdr_cells[4].text = 'Categorization'
                hdr_cells[5].text = 'Confidence in %'

                for username, created_on, target_type, post_url, categorization, confidence in paksitan_news:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(username)
                    row_cells[1].text = str(target_type)
                    row_cells[2].text = str(created_on)
                    row_cells[3].text = str(post_url)
                    row_cells[4].text = str(categorization)
                    row_cells[5].text = str(confidence)

            # --------------------------------- International News -------------------------------------

            if len(world_news) > 0:
                document.add_heading('World Wide News')
                table = document.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = main
                hdr_cells[1].text = 'URL'
                hdr_cells[2].text = 'Target Type'
                hdr_cells[3].text = 'Date'
                hdr_cells[4].text = 'Categorization'
                hdr_cells[5].text = 'Confidence in %'

                for username, created_on, target_type, post_url, categorization, confidence in world_news:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(username)
                    row_cells[1].text = str(target_type)
                    row_cells[2].text = str(created_on)
                    row_cells[3].text = str(post_url)
                    row_cells[4].text = str(categorization)
                    row_cells[5].text = str(confidence)

            # ---------------------------------------- BLOGS ------------------------------------
            if len(blogs) > 0:
                document.add_heading('Other websites')
                table = document.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = main
                hdr_cells[1].text = 'URL'
                hdr_cells[2].text = 'Target Type'
                hdr_cells[3].text = 'Date'
                hdr_cells[4].text = 'Categorization'
                hdr_cells[5].text = 'Confidence in %'

                for username, created_on, target_type, post_url, categorization, confidence in blogs:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(username)
                    row_cells[1].text = str(target_type)
                    row_cells[2].text = str(created_on)
                    row_cells[3].text = str(post_url)
                    row_cells[4].text = str(categorization)
                    row_cells[5].text = str(confidence)

            else:
                document.add_paragraph('No Data')

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=Categorized_Data.docx'

            document.save(response)
            return response


class LogView(viewsets.ViewSet):

    authentication_classes = (TokenAuthentication,)
    paginate_by = 10
    def get_user_log(self, request):
        auth_user = request.user.has_perm("core_management.tools_permission")
        print(auth_user)
        if auth_user == False:
            response = {
                'message': 'permission missing',
                'status': False,
                'result': None
                }
            return Response(response,status=status.HTTP_403_FORBIDDEN)
        log = Log.objects.filter(request_username=request.user)
        unique_user =Log.objects.order_by('request_username').values_list('request_username', flat=True).distinct('request_username')
        paginator = Paginator(log, 5000)
        log1 = paginator.page(1)
        all_data = []
        for each_log in log1:
            # print(str(each_log.request_url).split('/')[2])
            action = log_filter_url(str(each_log.request_url), each_log.request_method)
            if action:
                data = {}
                # print(action)
                data['user'] = str(each_log.request_username)
                data['message'] = action
                data['time'] = each_log.request_time
                all_data.append(data)
        # serializers = LogSerializers(log1, many=True)
        return Response({"user":unique_user,"history":all_data})

class LogHistoryView(viewsets.ModelViewSet):
    queryset = Log_history.objects.all()
    serializer_class = LogHistorySerializers
    authentication_classes = (TokenAuthentication,)


class ImageAnalysisFRSViewSet(viewsets.ModelViewSet):
    """
    ImageAnalysisFRS face recognetions 
    """
    queryset = ImageAnalysisFRS.objects.all()
    serializer_class = ImageAnalysisFRSSerializer
    authentication_classes = (TokenAuthentication,)
    # filter_fields = ('title',)
    # filter_backends = [filters.SearchFilter]
    
    # filter_fields = ('title', 'target_type')
    # filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    # filterset_fields = ['title', 'target_type']
    # search_fields = ['^title','^target_type']
    

    def list(self, request, *args, **kwargs):
        check = 0
        target_type = self.request.query_params.get('target_type')
        title = self.request.query_params.get('title')
        date = self.request.query_params.get('date')

        if title or target_type or date:
            if title:
                query = Q(title__icontains=title)
                check = check + 1

            if target_type:
                data = target_type
                data = data.replace("[","").replace("]","").replace("\'","").replace("\"","").split(",")
                for each_target in data:
                    print(each_target)
                    if check == 0:
                        check = check + 1
                        query = Q(target_type=each_target)
                    else:
                        query.add(Q(target_type=each_target), Q.OR)
            if date:
                if check == 0 :
                    query = Q(created_on__lt= date)
                else:
                    query.add(Q(created_on__lt= date), Q.AND)

            
            queryset = ImageAnalysisFRS.objects.filter(query).order_by('-created_on')
            serializer_class = ImageAnalysisFRSSerializer(queryset, many=True)
            return Response({"results":serializer_class.data})

        queryset = ImageAnalysisFRS.objects.all().order_by('-created_on')
        serializer_class = ImageAnalysisFRSSerializer(queryset, many=True)
        return Response({"results":serializer_class.data})

    def create(self, request, *args, **kwargs):
        try:
            file =  request.data['image']
            print(file)
            print(type(file))
            data = request.data.copy()
            # print(data)
            # print("-------after---------")

            if ImageAnalysisFRS.objects.filter(title= data["title"]).exists():
                return Response ({"message":"title '{0}' already present. title should be unique".format(data['title'])})


            data["image_url"] = ""
            test = perform_upload_frs(data, 'image')
            target_type, target_sub_type  = get_attribute(target_type=test['target_type'], sub_target_type=test['target_subtype'] )
            test['target_type'] = target_type
            test['target_subtype'] = target_sub_type
            serializer = self.get_serializer(data=test)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            bds_res = bds_object.image_analytics(test['image_url'], target_type, target_sub_type, test['title'])
            headers = self.get_success_headers(serializer.data)
            return Response({'message':"image uploaded Successfully","result":bds_res}, status=status.HTTP_201_CREATED, headers=headers)
    
        except Exception as e:
            print("exceptions ",e)
            return Response({'message':"exceptions:  {0}".format(e)})

    def destroy(self, request, *args, **kwargs):
        try:
            instance = self.get_object()
            self.perform_destroy(instance)

            # delete from ES
            es_object.delete_image_analysis_FRS(title=instance.title)
        except :
            pass
        return Response({"message":"Deleted Successfully"}, status=status.HTTP_204_NO_CONTENT)

# class Dashboard(viewsets.ViewSet):
#     authentication_classes = (TokenAuthentication,)
#     def dashboard_list(self, request):
#         return Response("its working..!")
class Dashboards(viewsets.ViewSet):

    authentication_classes = (TokenAuthentication,)

    def superuser_dashboard(self, request):

        social_media = ['fb', 'tw', 'rd', 'ln', 'yt',"in"]
        overview = {}
        
        final = []
        new_graph = {}
        round_graph = []
        all_activity = []
        all_targets_count = {}

        users_count = User.objects.all().count()
        case_count = Case.objects.all().count()
        social_target_count = SocialTarget.objects.filter(target_status_dic__6__message__icontains="Data Transformed").count()
        individual_count = Individual.objects.all().count()
        group_count = Group.objects.all().count()
        event_count = Event.objects.all().count()
        portfolio_count = individual_count + group_count + event_count
        # keybase_count = Keybase.objects.all().count()
        keybase_targets_count = KeybaseTarget.objects.filter(target_status_dic__6__message__icontains="Data Transformed").count()
        generic_targest_count = GenericTarget.objects.filter(target_status_dic__6__message__icontains="Data Transformed").count()

        # (AND: ('target_status_dic__6__message__icontains', 'Data Transformed'))

        # total = 0
        # for sm in social_media:
        #     count = 0
        #     social_data = SocialTarget.objects.filter(target_type=sm)
        #     for social_each_data in social_data:
        #         social_splitted_data=social_each_data.target_status_string.split(',')
        #         if 'Successful! Target Transformed successfully' in social_splitted_data:
        #             count = count + 1
        #     total = total + count
        #     final.append({"category":str(sm),"value":count})

        total = social_target_count  + generic_targest_count
        overview['users'] = users_count
        overview['cases'] = case_count
        overview['targets'] = total
        overview['portfolio'] = portfolio_count
        overview['keybase'] = keybase_targets_count
        all_cases = Case.objects.all()
        serialized_cases = CaseDashboardsSerializer(all_cases, many=True)

        all_keybase = KeybaseTarget.objects.filter(target_status_dic__6__message__icontains="Data Transformed")
        serialized_keybase = KeybaseTargetListSerializer(all_keybase, many=True)
        
        # keybase_dict={}
        # query = """ 
        #     SELECT target_type, count(*) as total
        #     FROM "target_management_socialtarget"
        #     WHERE "target_status_string" 
        #     LIKE '%Successful! Target Transformed successfully%'
        #     GROUP BY target_type
        
        # """
        # for sm_count in SocialTarget.objects.raw(query):
        #     final.append({"category":sm_count.target_type, 'value':sm_count.total})

        
        for sm in social_media:
            targets_on_social_media_list=[]
            targets_on_social_media = {}
            social_data = SocialTarget.objects.filter(target_type=sm)
            for social_each_data in social_data:
                
                social_splitted_data=social_each_data.target_status_string.split(',')
                for data_splited in social_splitted_data:
                    targets_social={}
                    if data_splited=='Successful! Target Transformed successfully':
                        targets_social['target_status_string']=social_each_data.target_status_string
                        targets_social['id']=social_each_data.id
                        targets_on_social_media_list.append(targets_social)
                    else:
                        pass   
            targets_on_social_media['value'] = len(targets_on_social_media_list)
            targets_on_social_media['targets']=targets_on_social_media_list
            targets_on_social_media['category'] = dict(INDEX_PLATFORM_ALL)[sm].split("_")[0]
            final.append(dict(targets_on_social_media))

        


        new_graph['category'] = 'keybase'
        new_graph['value'] = keybase_targets_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'generic'
        new_graph['value'] = generic_targest_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'social'
        new_graph['value'] = social_target_count
        round_graph.append(dict(new_graph))

        # users = User.objects.all()

       
        #team TMO results
        # teams = Team.objects.all()
        # team_list  = []
        # for team in teams:
        #     print("____ team  ____")
        #     team_dic = {}
        #     team_dic['team_name'] = str(team.team_name)
        #     team_memberes = Team_members.objects.filter(team=team)
        #     for each_team_member in team_memberes:
        #         u = each_team_member.user

                
        #         social_targets_count = SocialTarget.objects.filter(user_id=u.id).count()
        #         generic_targets_count = GenericTarget.objects.filter(user_id=u.id).count()
        #         keybase_targets_count = KeybaseTarget.objects.filter(user_id=u.id).count()

        #         # groups = Group.objects.filter(user=u).values()

                

        #         all_targets = social_targets_count + generic_targets_count + keybase_targets_count

                

        #         individual_count = Individual.objects.filter(user_id=u.id).count()
        #         group_count = Group.objects.filter(user_id=u.id).count()
        #         event_count = Event.objects.filter(user_id=u.id).count()
        #         linkanalysis_count = Log.objects.filter(request_username=u.id, request_url__icontains = '/v1/target/get/linkanalysis_explore/').count()

        #         portfolio_count = individual_count + group_count + event_count

                

        #         case_count = Case.objects.filter(user_id=u.id).count()
                

        #         keybase_targets_count = Keybase.objects.filter(user_id=u.id).count()
                
        #         team_dic['team_member_name'] = u.username
        #         team_dic['case'] = case_count
        #         team_dic['targets'] = all_targets
        #         team_dic['portfolio'] = portfolio_count
        #         team_dic['keybase'] = keybase_targets_count
        #         team_dic['link_analysis'] = linkanalysis_count

        #         team_list.append(dict(team_dic))
        
        teams = Team.objects.all()
        team_list  = []
        for team in teams:
            all_team_user_list  = []
            all_team_user_list.append(team.team_leader)
            print("____ team____")
            team_dic = {}
            team_dic['team_name'] = str(team.team_name)
            team_memberes = Team_members.objects.filter(team=team)
            for each_team_member in team_memberes:
                u = each_team_member.user
                all_team_user_list.append(u)
            social_targets_count_team = SocialTarget.objects.filter(user__in=all_team_user_list).count()
            generic_targets_count_team = GenericTarget.objects.filter(user__in=all_team_user_list).count()
            keybase_targets_count_team = KeybaseTarget.objects.filter(user__in=all_team_user_list).count()

            # groups = Group.objects.filter(user=u).values()

            

            all_targets = social_target_count + generic_targest_count + keybase_targets_count

            

            individual_count = Individual.objects.filter(user__in=all_team_user_list).count()
            group_count = Group.objects.filter(user__in=all_team_user_list).count()
            event_count = Event.objects.filter(user__in=all_team_user_list).count()
            linkanalysis_count = Log.objects.filter(request_username__in=all_team_user_list, request_url__icontains = '/v1/target/get/linkanalysis_explore/').count()

            portfolio_count = individual_count + group_count + event_count

            

            case_count = Case.objects.filter(user__in=all_team_user_list).count()
            

            keybase_targets_count = Keybase.objects.filter(user__in=all_team_user_list).count()
            
            team_dic['team_member_name'] = u.username
            team_dic['case'] = case_count
            team_dic['targets'] = all_targets
            team_dic['portfolio'] = portfolio_count
            team_dic['keybase'] = keybase_targets_count
            team_dic['link_analysis'] = linkanalysis_count

            team_list.append(dict(team_dic))

        
        # team_list  = []
        # for team in teams:
        #     print("____ team  ____")
        #     team_dic = {}
        #     u = team.team_leader
        #     team_memberes = Team_members.objects.filter(team=team)

        #     for each_team_member in team_memberes:
        #         print(each_team_member)
        #         u = each_team_member.user
        #         social_targets_count = SocialTarget.objects.filter(user_id=u.id).count()
        #         generic_targets_count = GenericTarget.objects.filter(user_id=u.id).count()
        #         keybase_targets_count = KeybaseTarget.objects.filter(user_id=u.id).count()

        #         # groups = Group.objects.filter(user=u).values()

                

        #         all_targets = social_targets_count + generic_targets_count + keybase_targets_count

                

        #         individual_count = Individual.objects.filter(user_id=u.id).count()
        #         group_count = Group.objects.filter(user_id=u.id).count()
        #         event_count = Event.objects.filter(user_id=u.id).count()
        #         linkanalysis_count = Log.objects.filter(request_username=u.id, request_url__icontains = '/v1/target/get/linkanalysis_explore/').count()

        #         portfolio_count = individual_count + group_count + event_count

                

        #         case_count = Case.objects.filter(user_id=u.id).count()
                

        #         keybase_targets_count = Keybase.objects.filter(user_id=u.id).count()
                
        #         team_dic['team_member_name'] = u.username
        #         team_dic['case'] = case_count
        #         team_dic['targets'] = all_targets
        #         team_dic['portfolio'] = portfolio_count
        #         team_dic['keybase'] = keybase_targets_count
        #         team_dic['link_analysis'] = linkanalysis_count

        #         team_list.append(dict(team_dic))



        response = {
            'overview': overview,
            'cms': serialized_cases.data,
            'kms': serialized_keybase.data,
            'targets_on_social_media': final,
            'new_graph': round_graph,
            # 'team_stats': all_activity, 
            "team_sta":team_list
        }
        return Response(response)
    def TMO_dashboard(self, request):

        if request.user.groups.filter(name='TMO').exists() == False:
            return Response("sorry you are not TMO")


        social_media = ['fb', 'tw', 'rd', 'ln', 'yt']
        overview = {}
        targets_on_social_media = {}
        final = []
        new_graph = {}
        round_graph = []
        all_activity = []
        all_targets_count = {}

        tmo = request.user
        print(tmo)
        team = Team.objects.filter(team_leader=tmo).first()

        print(team)
        team_memberes = Team_members.objects.filter(team=team)
        print(team_memberes)
        list_of_user = []
        for each_team_member in team_memberes:
            list_of_user.append(each_team_member.user)
        list_of_user.append(tmo)
        print('---------list of user------------', list_of_user)

        # users_count = User.objects.all().count()
        case_count = Case.objects.filter(user__in=list_of_user).count()
        target_count = SocialTarget.objects.filter(user__in=list_of_user).count()
        individual_count = Individual.objects.filter(user__in=list_of_user).count()
        group_count = Group.objects.filter(user__in=list_of_user).count()
        event_count = Event.objects.filter(user__in=list_of_user).count()
        portfolio_count = individual_count + group_count + event_count
        keybase_count = Keybase.objects.filter(user__in=list_of_user).count()
        keybase_targets_count = KeybaseTarget.objects.filter(user__in=list_of_user).count()
        generic_targest_count = GenericTarget.objects.filter(user__in=list_of_user).count()

        # overview['users'] = users_count
        overview['cases'] = case_count
        overview['targets'] = target_count
        overview['portfolio'] = portfolio_count
        overview['keybase'] = keybase_targets_count
        all_cases = Case.objects.filter(user__in=list_of_user)
        serialized_cases = CaseDashboardsSerializer(all_cases, many=True)

        all_keybase = KeybaseTarget.objects.filter(user__in=list_of_user)
        serialized_keybase = KeybaseTargetListSerializer(all_keybase, many=True)

        
        targets_on_social_media_list=[]
        for sm in social_media:
            social_data = SocialTarget.objects.filter(target_type=sm)
            for social_each_data in social_data:
                targets_social={}
                social_splitted_data=social_each_data.target_status_string.split(',')
                for data_splited in social_splitted_data:
                    if data_splited=='Successful! Target Transformed successfully':
                        targets_social['target_status_string']=social_each_data.target_status_string
                        targets_social['id']=social_each_data.id
                        targets_on_social_media_list.append(targets_social)
                    else:
                        pass         
            targets_on_social_media['targets']=targets_on_social_media_list
            targets_on_social_media['category'] = dict(INDEX_PLATFORM_ALL)[sm].split("_")[0] 
            targets_on_social_media['value'] = len(targets_on_social_media_list)
            final.append(dict(targets_on_social_media))

        new_graph['category'] = 'keybase'
        new_graph['value'] = keybase_targets_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'generic'
        new_graph['value'] = generic_targest_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'social'
        new_graph['value'] = final
        round_graph.append(dict(new_graph))

 
        # teams = Team.objects.all()
        team_list  = []
        # team_memberes = Team_members.objects.filter(team=team)
        for each_team_member in team_memberes:
            u = each_team_member.user
            team_dic = {}
            if u.groups.filter(name='AMO').exists():
                print("AMO found {}".format(u.username))
                

            

            
            # u = team.team_leader
            social_targets_count = SocialTarget.objects.filter(user_id=u.id).count()
            generic_targets_count = GenericTarget.objects.filter(user_id=u.id).count()
            keybase_targets_count = KeybaseTarget.objects.filter(user_id=u.id).count()

            # groups = Group.objects.filter(user=u).values()

            

            all_targets = social_targets_count + generic_targets_count + keybase_targets_count

            

            individual_count = Individual.objects.filter(user_id=u.id).count()
            group_count = Group.objects.filter(user_id=u.id).count()
            event_count = Event.objects.filter(user_id=u.id).count()
            linkanalysis_count = Log.objects.filter(request_username=u.id, request_url__icontains = '/v1/target/get/linkanalysis_explore/').count()

            portfolio_count = individual_count + group_count + event_count

            case_count = Case.objects.filter(user_id=u.id).count()
            keybase_targets_count = Keybase.objects.filter(user_id=u.id).count()
            
            team_dic['group_team_member_name'] = u.groups.first().name + " username " + u.username
            team_dic['case'] = case_count
            team_dic['targets'] = all_targets
            team_dic['portfolio'] = portfolio_count
            team_dic['keybase'] = keybase_targets_count
            team_dic['link_analysis'] = linkanalysis_count

            team_list.append(dict(team_dic))


        response = {
            'overview': overview,
            'cms': serialized_cases.data,
            'kms': serialized_keybase.data,
            'targets_on_social_media': final,
            'new_graph': round_graph,
            # 'team_stats': all_activity, 
            "team_sta":team_list
        }
        return Response(response)


    def simple_dashboard(self, request):

        if request.user.groups.filter(~Q(name="TMO")).exists() == False:
            return Response("sorry you are not TAO")


        social_media = ['fb', 'tw', 'rd', 'ln', 'rd', 'yt']
        overview = {}
        targets_on_social_media = {}
        final = []
        new_graph = {}
        round_graph = []
        all_activity = []
        all_targets_count = {}
        team_list = []

        user = request.user



        # users_count = User.objects.all().count()
        case_count = Case.objects.filter(user=user).count()
        target_count = SocialTarget.objects.filter(user=user).count()
        individual_count = Individual.objects.filter(user=user).count()
        group_count = Group.objects.filter(user=user).count()
        event_count = Event.objects.filter(user=user).count()
        portfolio_count = individual_count + group_count + event_count
        keybase_count = Keybase.objects.filter(user=user).count()
        keybase_targets_count = KeybaseTarget.objects.filter(user=user).count()
        generic_targest_count = GenericTarget.objects.filter(user=user).count()

        # overview['users'] = users_count
        overview['cases'] = case_count
        overview['targets'] = target_count
        overview['portfolio'] = portfolio_count
        overview['keybase'] = keybase_targets_count
        all_cases = Case.objects.filter(user=user)
        serialized_cases = CaseDashboardsSerializer(all_cases, many=True)

        all_keybase = KeybaseTarget.objects.filter(user=user)
        serialized_keybase = KeybaseTargetListSerializer(all_keybase, many=True)

        
        targets_on_social_media_list=[]
        for sm in social_media:
            social_data = SocialTarget.objects.filter(target_type=sm)
            for social_each_data in social_data:
                targets_social={}
                social_splitted_data=social_each_data.target_status_string.split(',')
                for data_splited in social_splitted_data:
                    if data_splited=='Successful! Target Transformed successfully':
                        targets_social['target_status_string']=social_each_data.target_status_string
                        targets_social['id']=social_each_data.id
                        targets_on_social_media_list.append(targets_social)
                    else:
                        pass         
            targets_on_social_media['targets']=targets_on_social_media_list
            targets_on_social_media['category'] = sm
            targets_on_social_media['value'] = len(targets_on_social_media_list)
            final.append(dict(targets_on_social_media))

        new_graph['category'] = 'keybase'
        new_graph['value'] = keybase_targets_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'generic'
        new_graph['value'] = generic_targest_count
        round_graph.append(dict(new_graph))
        new_graph['category'] = 'social'
        new_graph['value'] = final
        round_graph.append(dict(new_graph))

 

        social_targets_count = SocialTarget.objects.filter(user=user).count()
        generic_targets_count = GenericTarget.objects.filter(user=user).count()
        keybase_targets_count = KeybaseTarget.objects.filter(user=user).count()


        

        all_targets = social_targets_count + generic_targets_count + keybase_targets_count

        
        team_dic = {}
        individual_count = Individual.objects.filter(user=user).count()
        group_count = Group.objects.filter(user=user).count()
        event_count = Event.objects.filter(user=user).count()
        linkanalysis_count = Log.objects.filter(request_username=user, request_url__icontains = '/v1/target/get/linkanalysis_explore/').count()

        portfolio_count = individual_count + group_count + event_count

        case_count = Case.objects.filter(user=user).count()
        keybase_targets_count = Keybase.objects.filter(user=user).count()
        
        team_dic['group_team_member_name'] = user.groups.first().name + " username " + user.username
        team_dic['case'] = case_count
        team_dic['targets'] = all_targets
        team_dic['portfolio'] = portfolio_count
        team_dic['keybase'] = keybase_targets_count
        team_dic['link_analysis'] = linkanalysis_count

        team_list.append(dict(team_dic))


        response = {
            'overview': overview,
            'cms': serialized_cases.data,
            'kms': serialized_keybase.data,
            'targets_on_social_media': final,
            'new_graph': round_graph,
            "team_sta":team_list
        }
        return Response(response)

class Sftp(viewsets.ViewSet):
    authentication_classes = (TokenAuthentication,)
    
    def sftp_delete(self, request, *args, **kwargs):
        # return Response({"message":"sftp_deleted"})
        gtr = kwargs['gtr']
        print(gtr)
        ftp = FTP_Client()
        ftp.login()
        check = ftp.delete_dir_file(gtr)

        # if True:
        return Response({"message":"Successfully deleted",
                                "result": str(check)})
                                
